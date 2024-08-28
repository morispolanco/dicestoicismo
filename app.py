import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Filosófico del Estoicismo", page_icon="📚", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Filosófico basado en el pensamiento de la Filosofía Estoica. Permite a los usuarios obtener artículos de diccionario o ensayos académicos sobre términos filosóficos, incluyendo discusiones de varios autores estoicos con referencias a sus obras.

    ### Cómo usar la aplicación:

    1. Elija un término filosófico de la lista predefinida o proponga su propio término.
    2. Seleccione el tipo de contenido que desea generar: artículo de diccionario o ensayo académico.
    3. Haga clic en "Generar contenido" para crear el artículo o ensayo.
    4. Lea el contenido generado y las fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 27 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Filosófico del Estoicismo* [Aplicación web]. https://estoicismo.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar contenido basado en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario Filosófico del Estoicismo")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # 101 philosophical terms related to Stoicism
    terminos_filosoficos = sorted([
        "Apatía", "Apatheia", "Ataraxia", "Autarquía", "Bien", "Catalepsis", "Cosmopolitismo", "Deber",
        "Destino", "Dikaiosyne", "Disciplines of Assent", "Disciplines of Desire", "Disciplines of Action",
        "Ekpyrosis", "Eudaimonia", "Felicidad", "Física", "Hegemonikon", "Impresiones", "Indiferentes",
        "Juicio", "Kathêkon", "Logos", "Lógica", "Mal", "Memoria", "Moderación", "Naturaleza", "Oikeiôsis",
        "Pasiones", "Phantasia", "Pneuma", "Prohairesis", "Prokopton", "Prólepsis", "Providencia",
        "Prudencia", "Razón", "Sabio", "Sympatheia", "Temperancia", "Virtud", "Vivir conforme a la naturaleza",
        "Voluntad", "Amor fati", "Ascesis", "Asentimiento", "Autodominio", "Bien común", "Constancia",
        "Contemplación", "Coraje", "Decoro", "Dike", "Disciplina", "Dualismo", "Ecuanimidad", "Ejemplaridad",
        "Ejercicios espirituales", "Epoche", "Esencia", "Espectador desapasionado", "Estoicismo", "Ética",
        "Examen de conciencia", "Fatalismo", "Fortaleza", "Fraternidad universal", "Imperturbabilidad",
        "Impermanencia", "Indiferencia", "Juicio de valor", "Justicia", "Libertad interior", "Meditación",
        "Memento mori", "Moderación", "Moralidad", "Muerte", "Naturaleza humana", "Necesidad", "Panteísmo",
        "Pasiones", "Physis", "Piedad", "Placer", "Pneuma", "Prenoción", "Progreso moral", "Providencia",
        "Racionalidad", "Resignación", "Sabiduría", "Serenidad", "Sophrosyne", "Superación personal", "Telos",
        "Theoria", "Tranquilidad", "Valor"
    ])

    def buscar_informacion(query):
        url = "https://api.serply.io/v1/scholar"
        params = {
            "q": f"{query} Filosofía Estoica"
        }
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.get(url, headers=headers, params=params)
        return response.json()

    def generar_contenido(termino, tipo_contenido, contexto):
        url = "https://api.together.xyz/inference"
        if tipo_contenido == "Generar artículo de diccionario":
            instruccion = f"Proporciona un artículo de diccionario filosófico sobre el término '{termino}' en el contexto del estoicismo. Incluye definiciones y perspectivas de varios autores estoicos importantes, citando sus obras específicas. El artículo debe ser conciso pero informativo, similar a una entrada de enciclopedia filosófica."
        else:  # Generar ensayo académico
            instruccion = f"Escribe un ensayo académico sobre el concepto de '{termino}' en la filosofía estoica. El ensayo debe incluir una introducción, desarrollo y conclusión. Analiza el término desde la perspectiva de varios autores estoicos importantes, citando sus obras específicas. Discute la evolución del concepto y su relevancia en el pensamiento estoico. Concluye con una reflexión sobre la importancia de este concepto en la filosofía estoica y su posible aplicación en la actualidad."

        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nInstrucción: {instruccion}\n\nContenido:",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1.1,
            "stop": ["Instrucción:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, contenido, tipo_contenido, fuentes):
        doc = Document()
        doc.add_heading(f'Diccionario Filosófico del Estoicismo - {tipo_contenido}', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading(tipo_contenido, level=1)
        doc.add_paragraph(contenido)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término filosófico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_filosoficos)
    else:
        termino = st.text_input("Ingresa tu propio término filosófico:")

    tipo_contenido = st.radio("Selecciona el tipo de contenido a generar:", ["Generar artículo de diccionario", "Generar ensayo académico"])

    if st.button("Generar contenido"):
        if termino:
            with st.spinner("Buscando información y generando contenido..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                contexto = "\n".join([item.get("snippet", "") for item in resultados_busqueda.get("organic", [])])
                fuentes = [item.get("link", "") for item in resultados_busqueda.get("organic", [])]

                # Generar contenido
                contenido = generar_contenido(termino, tipo_contenido, contexto)

                # Mostrar el contenido
                st.subheader(f"{tipo_contenido} para el término: {termino}")
                st.markdown(contenido)

                # Botón para descargar el documento
                doc = create_docx(termino, contenido, tipo_contenido, fuentes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar contenido en DOCX",
                    data=buffer,
                    file_name=f"{tipo_contenido.split()[-1]}_{termino.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona o ingresa un término.")
